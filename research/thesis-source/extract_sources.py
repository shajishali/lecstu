from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parent


def extract_docx(source: Path, target: Path) -> None:
    document = Document(source)
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            style = paragraph.style.name if paragraph.style else ""
            lines.append(f"[{style}] {text}")
    for table_number, table in enumerate(document.tables, start=1):
        lines.append(f"\n[TABLE {table_number}]")
        for row in table.rows:
            cells = [cell.text.replace("\n", " / ").strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    target.write_text("\n".join(lines), encoding="utf-8")


def extract_pdf(source: Path, target: Path) -> None:
    reader = PdfReader(source)
    pages: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        pages.append(f"\n\n===== PAGE {page_number} =====\n\n{page.extract_text() or ''}")
    target.write_text("".join(pages), encoding="utf-8")


extract_docx(ROOT / "thesis-template.docx", ROOT / "thesis-template.txt")
extract_pdf(ROOT / "proposal.pdf", ROOT / "proposal.txt")
extract_pdf(ROOT / "senior-sample.pdf", ROOT / "senior-sample.txt")
